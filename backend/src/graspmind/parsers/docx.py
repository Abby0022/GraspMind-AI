"""DOCX document parser using python-docx.

Extracts text from Word documents while preserving paragraph
structure and heading hierarchy.
"""

import logging
from pathlib import Path

from docx import Document

from graspmind.parsers.pdf import ParsedDocument, ParsedPage

logger = logging.getLogger(__name__)


def parse_docx(file_path: str | Path) -> ParsedDocument:
    """Parse a DOCX file and extract text content.

    Paragraphs are grouped into logical pages of ~3000 chars
    since DOCX files don't have physical pages.

    Args:
        file_path: Path to the DOCX file.

    Returns:
        ParsedDocument with extracted text.
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"DOCX not found: {path}")

    try:
        doc = Document(str(path))
    except Exception as exc:
        raise ValueError(f"Failed to open DOCX: {exc}") from exc

    # Extract all paragraphs with style info
    paragraphs: list[dict] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        style_name = para.style.name if para.style else ""
        is_heading = style_name.startswith("Heading")
        heading_level = 0
        if is_heading:
            try:
                heading_level = int(style_name.replace("Heading", "").strip())
            except ValueError:
                heading_level = 1

        paragraphs.append({
            "text": text,
            "style": style_name,
            "is_heading": is_heading,
            "heading_level": heading_level,
        })

    # Also extract text from tables
    for table in doc.tables:
        table_rows: list[str] = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                table_rows.append(" | ".join(cells))
        if table_rows:
            paragraphs.append({
                "text": "\n".join(table_rows),
                "style": "Table",
                "is_heading": False,
                "heading_level": 0,
            })

    # Group paragraphs into logical "pages" (~3000 chars each)
    pages: list[ParsedPage] = []
    current_text: list[str] = []
    current_headings: list[str] = []
    current_chars = 0
    page_num = 1
    chars_per_page = 3000

    for para in paragraphs:
        if para["is_heading"]:
            current_headings.append(para["text"])

        current_text.append(para["text"])
        current_chars += len(para["text"])

        # Split into new page on heading boundaries or char limit
        if current_chars >= chars_per_page and para["is_heading"]:
            # Save current page (without the heading that starts the next)
            page_text = "\n\n".join(current_text[:-1])
            if page_text.strip():
                pages.append(ParsedPage(
                    page_number=page_num,
                    text=page_text,
                    headings=current_headings[:-1],
                ))
                page_num += 1

            current_text = [para["text"]]
            current_headings = [para["text"]]
            current_chars = len(para["text"])

    # Don't forget the last page
    if current_text:
        pages.append(ParsedPage(
            page_number=page_num,
            text="\n\n".join(current_text),
            headings=current_headings,
        ))

    # Title: first heading or filename
    title = path.stem
    headings = [p for p in paragraphs if p["is_heading"]]
    if headings:
        title = headings[0]["text"]

    # Get core properties
    core_props = doc.core_properties
    if core_props and core_props.title:
        title = core_props.title

    logger.info("Parsed DOCX: %s (%d pages, %d paragraphs)", title, len(pages), len(paragraphs))

    return ParsedDocument(
        title=title,
        pages=pages,
        total_pages=len(pages),
        source_type="docx",
        metadata={
            "file_name": path.name,
            "file_size": path.stat().st_size,
            "paragraph_count": len(paragraphs),
        },
    )
